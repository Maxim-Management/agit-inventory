"""Генерация «Акта на списание материалов» (.docx) по форме, предоставленной
заказчиком (файл SPF019WOF.doc) — печатный документ для комиссии по итогам
ремонтной/сборочной работы.

Автоматически заполняются: дата акта (по дате работы), список списанных на
эту работу материалов (наименование детали + серийный номер для серийных,
количество, сумма, причина списания) и итоговая сумма. Состав комиссии
(должности/ФИО), подпись директора и сумма прописью оставлены пустыми для
заполнения от руки/печати — эти данные система не ведёт, а автоматическая
транслитерация суммы прописью для официального документа обычно требует
дополнительной ручной проверки, поэтому сознательно не подставляется."""
import io
from datetime import date

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

MONTHS_RU = [
    "января", "февраля", "марта", "апреля", "мая", "июня",
    "июля", "августа", "сентября", "октября", "ноября", "декабря",
]


def _parse_date(value):
    if value is None:
        return None
    if isinstance(value, date):
        return value
    s = str(value)[:10]
    try:
        y, m, d = s.split("-")
        return date(int(y), int(m), int(d))
    except Exception:
        return None


def _format_date_ru(value):
    """«25» августа 2026г. — как в шапке акта; при отсутствии даты —
    прочерки для заполнения от руки, как в исходной форме."""
    d = _parse_date(value)
    if d is None:
        return "«___» ______________ 20__г."
    return f"«{d.day:02d}» {MONTHS_RU[d.month - 1]} {d.year}г."


def _fmt_money(value):
    """Та же типографика, что и в веб-интерфейсе (app/__init__.py: fmt_num) —
    неразрывный пробел между разрядами, запятая перед копейками."""
    try:
        f = float(value or 0)
    except (TypeError, ValueError):
        f = 0.0
    sign = "-" if f < 0 else ""
    f = abs(f)
    grouped = f"{f:,.2f}"
    int_part, _, dec_part = grouped.partition(".")
    int_part = int_part.replace(",", " ")
    return f"{sign}{int_part},{dec_part}"


def _set_col_widths(table, widths):
    """widths — список длин python-docx (например, Pt(...)) по числу столбцов."""
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w


def build_writeoff_act_docx(job, write_off_rows, reason_labels, city="г. Иркутск"):
    """job — строка service_jobs (+ tool_serial); write_off_rows — результат
    app/jobs.py: job_write_offs() (уже содержит unit_cost_rub/line_cost_rub);
    reason_labels — REASON_LABELS из app/analytics.py."""
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Pt(85)
    section.right_margin = Pt(56)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # --- «Утверждаю: Директор ...» — правый верхний угол, как в форме ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("Утверждаю:").bold = True
    p2 = doc.add_paragraph("Директор")
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p3 = doc.add_paragraph("_________________Рябов М.И.")
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p4 = doc.add_paragraph('"__"____________20_г.')
    p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("АКТ НА СПИСАНИЕ МАТЕРИАЛОВ")
    run.bold = True
    run.font.size = Pt(14)

    dateline = doc.add_paragraph()
    dateline.add_run(f"{city}\t\t\t\t\t\t\t\t{_format_date_ru(job.get('job_date'))}")

    doc.add_paragraph()
    doc.add_paragraph("Комиссия в составе:")

    # --- таблица состава комиссии — пустая, заполняется от руки, как в исходной форме ---
    commission = doc.add_table(rows=6, cols=3)
    commission.style = "Table Grid"
    commission.alignment = WD_TABLE_ALIGNMENT.CENTER
    for member_row in (0, 2, 4):
        for c in range(3):
            commission.cell(member_row, c).text = ""
        label_row = member_row + 1
        for c, label in enumerate(("Должность", "ФИО", "Подпись")):
            commission.cell(label_row, c).text = label

    doc.add_paragraph()
    doc.add_paragraph(
        f"Произвела списание материалов в связи с нижеуказанными причинами "
        f"по работе «{job.get('title') or job.get('tool_assembly') or ('№ ' + str(job.get('id')))}»"
        f"{' на инструменте ' + job['tool_serial'] if job.get('tool_serial') else ''}:"
    )

    # --- таблица материалов, автоматически заполненная списаниями по работе ---
    materials = doc.add_table(rows=1, cols=5)
    materials.style = "Table Grid"
    header_cells = materials.rows[0].cells
    for cell, text in zip(header_cells, ("№ п/п", "Наименование материала", "Количество", "Сумма, ₽", "Причина списания")):
        cell.text = text
        for para in cell.paragraphs:
            for r in para.runs:
                r.bold = True

    total = 0.0
    for i, w in enumerate(write_off_rows, start=1):
        qty = float(w.get("quantity") or 0)
        line_cost = float(w.get("line_cost_rub") or 0)
        total += line_cost
        name = w.get("part_name") or ""
        if w.get("serial_number"):
            name += f" (с/н {w['serial_number']})"
        elif w.get("part_number"):
            name += f" ({w['part_number']})"
        row = materials.add_row().cells
        row[0].text = str(i)
        row[1].text = name
        row[2].text = f"{qty:g}"
        row[3].text = _fmt_money(line_cost)
        row[4].text = reason_labels.get(w.get("reason"), w.get("reason") or "")

    _set_col_widths(materials, [Pt(30), Pt(190), Pt(60), Pt(70), Pt(90)])

    doc.add_paragraph()
    total_p = doc.add_paragraph()
    total_p.add_run(f"Всего на сумму: {_fmt_money(total)} ₽ _____________________________(сумма прописью).")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
